
import streamlit as st
import pandas as pd
from docx import Document
from pathlib import Path
import re
import tempfile

st.header("📄 Generador de Contratos CITARA")

def safe_filename(s):
    return re.sub(r'[\\/:*?"<>|]+', '_', str(s))

def replace_tokens(paragraph, datos):
    full_text = ''.join(run.text for run in paragraph.runs)
    for k, v in datos.items():
        full_text = full_text.replace(f"[{k}]", v)
    for i in range(len(paragraph.runs)):
        paragraph.runs[i].text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text

uploaded_file = st.file_uploader("Sube tu Excel con datos de contratos", type=["xlsx"])
template_file = st.file_uploader("Sube la plantilla de contrato (.docx)", type=["docx"])

if uploaded_file and template_file:
    df = pd.read_excel(uploaded_file)
    st.success("✅ Archivos cargados correctamente")

    st.dataframe(df.head())

    if st.button("Generar contratos"):
        with st.spinner("Generando contratos..."):
            output_files = []
            out_dir = Path(tempfile.mkdtemp())

            for _, row in df.iterrows():
                datos = {col: str(row[col]) if pd.notna(row[col]) else '' for col in df.columns}
                doc = Document(template_file)

                for p in doc.paragraphs:
                    replace_tokens(p, datos)
                for table in doc.tables:
                    for row_ in table.rows:
                        for cell in row_.cells:
                            for p in cell.paragraphs:
                                replace_tokens(p, datos)

                filename = f"CONTRATO_{safe_filename(datos.get('PAC', ''))}_{safe_filename(datos.get('CENTRO', ''))}.docx"
                filepath = out_dir / filename
                doc.save(filepath)
                output_files.append(filepath)

            st.success(f"✅ {len(output_files)} contrato(s) generados.")

            for file in output_files:
                st.download_button(
                    label=f"Descargar {file.name}",
                    data=file.read_bytes(),
                    file_name=file.name
                )
